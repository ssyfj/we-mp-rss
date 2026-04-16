"""
PDF 内容提取工具

功能特性：
- 从 PDF 提取文本内容并转换为 Markdown
- 从 PDF 提取内容并转换为 DOCX
- 支持保留格式和布局

依赖库：
- markitdown: 用于 PDF 转 Markdown
- pymupdf: 用于 PDF 处理
- python-docx: 用于生成 DOCX

使用示例：
    from tools.mdtools.pdf_extractor import pdf_to_markdown, pdf_to_docx
    
    # PDF 转 Markdown
    markdown_content = pdf_to_markdown("input.pdf")
    
    # PDF 转 DOCX
    pdf_to_docx("input.pdf", "output.docx")
"""
import os
import logging
from typing import Optional
from pathlib import Path

try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError:
    MARKITDOWN_AVAILABLE = False
    logging.warning("markitdown 未安装，PDF 转 Markdown 功能将受限")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF 未安装，PDF 处理功能将受限")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx 未安装，PDF 转 DOCX 功能将受限")

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    logging.warning("pdf2docx 未安装，建议安装: pip install pdf2docx")


def pdf_to_markdown(pdf_path: str) -> str:
    """
    使用 markitdown 从 PDF 提取内容并转换为 Markdown
    
    Args:
        pdf_path: PDF 文件路径
    
    Returns:
        str: Markdown 内容
    """
    if not MARKITDOWN_AVAILABLE:
        raise ImportError("需要安装 markitdown: pip install markitdown")
    
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")
    
    try:
        # 使用 markitdown 转换
        md = MarkItDown()
        result = md.convert(pdf_path)
        return result.text_content
    except Exception as e:
        logging.error(f"PDF 转 Markdown 失败: {e}")
        raise


def pdf_to_docx(pdf_path: str, output_path: str) -> bool:
    """
    使用 pdf2docx 库从 PDF 转换为 DOCX（推荐方法）
    如果 pdf2docx 不可用，则回退到 PyMuPDF 方法
    
    Args:
        pdf_path: PDF 文件路径
        output_path: 输出 DOCX 文件路径
    
    Returns:
        bool: 是否成功
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")
    
    # 优先使用 pdf2docx 库（更好的转换质量）
    if PDF2DOCX_AVAILABLE:
        try:
            # 确保输出目录存在
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            cv = Converter(pdf_path)
            cv.convert(
                output_file, 
                layout=True,        # 1. 开启布局分析（最重要）
                tables=True,        # 2. 解析表格（默认开启）
                images=True,        # 3. 提取图片（默认开启）
                rotate=True,        # 4. 自动旋转页面（处理横向页面很有用）
                multi_processing=True # 5. 开启多进程加速
            )
            cv.close()
            
            logging.info(f"使用 pdf2docx 成功转换: {output_path}")
            return True
            
        except Exception as e:
            logging.error(f"pdf2docx 转换失败: {e}，尝试使用 PyMuPDF 方法")
            # 如果 pdf2docx 失败，继续尝试 PyMuPDF 方法
    
    # 回退到 PyMuPDF 方法
    if not PYMUPDF_AVAILABLE:
        raise ImportError("需要安装 PyMuPDF: pip install pymupdf 或 pdf2docx: pip install pdf2docx")
    
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    
    try:
        import io
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 打开 PDF
        pdf_doc = fitz.open(pdf_path)
        
        # 创建 Word 文档
        doc = Document()
        
        # 遍历每一页
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # 提取文本内容
            text = page.get_text()
            if text.strip():
                # 按段落分割文本
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        # 设置字体
                        for run in p.runs:
                            run.font.name = 'SimSun'
                            run.font.size = Pt(12)
            
            # 提取图片（使用 get_images 方法）
            image_list = page.get_images(full=True)
            if image_list:
                logging.info(f"页面 {page_num + 1} 找到 {len(image_list)} 张图片")
                
                for img_index, img_info in enumerate(image_list):
                    try:
                        # 获取图片的 xref
                        xref = img_info[0]
                        
                        # 提取图片数据
                        base_image = pdf_doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # 创建图片流
                        image_stream = io.BytesIO(image_bytes)
                        
                        # 获取图片尺寸信息
                        width = base_image.get("width", 0)
                        height = base_image.get("height", 0)
                        
                        # 如果没有尺寸信息，使用默认值
                        if width <= 0 or height <= 0:
                            width_inches = 4  # 默认宽度
                        else:
                            # 转换为英寸（假设 72 DPI）
                            width_inches = width / 72
                            
                            # 限制最大宽度为 6 英寸
                            max_width = 6
                            if width_inches > max_width:
                                width_inches = max_width
                        
                        # 添加段落并插入图片
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(image_stream, width=Inches(width_inches))
                        
                        logging.info(f"已添加图片 {img_index + 1}: 宽度 {width_inches:.2f} 英寸")
                    
                    except Exception as img_error:
                        logging.warning(f"提取图片 {img_index + 1} 失败: {img_error}")
                        continue
        
        # 确保输出目录存在
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # 保存文档
        doc.save(output_path)
        logging.info(f"使用 PyMuPDF 成功转换: {output_path}")
        return True
        
    except Exception as e:
        logging.error(f"PDF 转 DOCX 失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def pdf_to_text(pdf_path: str) -> str:
    """
    使用 PyMuPDF 从 PDF 提取纯文本
    
    Args:
        pdf_path: PDF 文件路径
    
    Returns:
        str: 纯文本内容
    """
    if not PYMUPDF_AVAILABLE:
        raise ImportError("需要安装 PyMuPDF: pip install pymupdf")
    
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")
    
    try:
        # 打开 PDF
        doc = fitz.open(pdf_path)
        text_content = []
        
        # 遍历每一页
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            text_content.append(text)
        
        return "\n\n".join(text_content)
        
    except Exception as e:
        logging.error(f"PDF 提取文本失败: {e}")
        raise


if __name__ == "__main__":
    # 测试代码
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python pdf_extractor.py <pdf_file>")
        sys.exit(1)
    
    pdf_file = sys.argv[1]
    
    try:
        # 测试 PDF 转 Markdown
        print("测试 PDF 转 Markdown...")
        markdown = pdf_to_markdown(pdf_file)
        print(f"Markdown 内容长度: {len(markdown)}")
        print("前 500 字符:")
        print(markdown[:500])
        
        # 测试 PDF 转 DOCX
        print("\n测试 PDF 转 DOCX...")
        output_docx = pdf_file.replace('.pdf', '.docx')
        success = pdf_to_docx(pdf_file, output_docx)
        if success:
            print(f"✓ DOCX 生成成功: {output_docx}")
        else:
            print("✗ DOCX 生成失败")
            
    except Exception as e:
        print(f"✗ 测试失败: {e}")
        import traceback
        traceback.print_exc()
