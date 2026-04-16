#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HTML to Word Document Converter Tool

This module provides a tool for converting HTML content to Word documents.
It supports various HTML elements including headers, paragraphs, lists, 
code blocks, tables, images, and more.

Usage Examples:
    # 从HTML字符串转换
    from tools.mdtools.html2docx import html_to_docx
    html_to_docx(html_content, 'output.docx')
    
    # 使用转换器类
    from tools.mdtools.html2docx import HtmlToDocxConverter
    converter = HtmlToDocxConverter()
    converter.convert(html_content, 'output.docx')

Author: AI Assistant
Date: 2025/04/15
"""

from hashlib import sha256
import os
import re
import logging
import tempfile
import requests
import time
import random
from typing import Optional, Dict, Any, List
from pathlib import Path
from urllib.parse import urlparse

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    raise ImportError("请安装 python-docx 库: pip install python-docx")

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    raise ImportError("请安装 beautifulsoup4 库: pip install beautifulsoup4")

try:
    from PIL import Image
except ImportError:
    raise ImportError("请安装 Pillow 库: pip install pillow")


class HtmlToDocxConverter:
    """
    HTML 转 Word 文档转换器
    
    支持的 HTML 元素：
    - 标题 (h1-h6)
    - 段落 (p)
    - 粗体和斜体 (strong, b, em, i)
    - 有序和无序列表 (ul, ol, li)
    - 代码块 (pre, code)
    - 表格 (table, tr, td, th)
    - 链接 (a)
    - 图片 (img)
    - 引用块 (blockquote)
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None, document_title: Optional[str] = None):
        """
        初始化转换器
        
        Args:
            config: 配置字典，包含样式和格式设置
            document_title: 文档标题，将插入到文档最上部分
        """
        self.config = self._get_default_config()
        if config:
            self.config.update(config)
        
        self.document_title = document_title
        self.document = None
        self.logger = self._setup_logger()
    
    def _get_default_config(self) -> Dict[str, Any]:
        """获取默认配置"""
        return {
            'font_name': '宋体',
            'font_size': 12,
            'line_spacing': 1.15,
            'paragraph_spacing_before': 6,
            'paragraph_spacing_after': 6,
            'heading_styles': {
                1: {'size': 18, 'bold': True, 'color': RGBColor(0, 0, 0)},
                2: {'size': 16, 'bold': True, 'color': RGBColor(0, 0, 0)},
                3: {'size': 14, 'bold': True, 'color': RGBColor(0, 0, 0)},
                4: {'size': 13, 'bold': True, 'color': RGBColor(0, 0, 0)},
                5: {'size': 12, 'bold': True, 'color': RGBColor(0, 0, 0)},
                6: {'size': 11, 'bold': True, 'color': RGBColor(0, 0, 0)},
            },
            'code_font': 'Consolas',
            'code_background': RGBColor(245, 245, 245),
            'table_style': 'Table Grid',
            'remove_links': False,  # 是否去除链接
            'remove_images': False,  # 是否去除图片
            'download_delay_min': 1.0,  # 图片下载最小延时（秒）
            'download_delay_max': 3.0,  # 图片下载最大延时（秒）
        }
    
    def _setup_logger(self) -> logging.Logger:
        """设置日志记录器"""
        logger = logging.getLogger(__name__)
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def convert(self, html_content: str, output_file: str, document_title: Optional[str] = None) -> bool:
        """
        转换 HTML 内容为 Word 文档
        
        Args:
            html_content: HTML 文本内容
            output_file: 输出的 Word 文件路径
            document_title: 文档标题，将插入到文档最上部分
            
        Returns:
            bool: 转换是否成功
        """
        try:
            # 设置文档标题（如果提供）
            if document_title:
                self.document_title = document_title
            
            # 创建新的 Word 文档
            self.document = Document()
            self._setup_document_styles()
            
            # 添加文档标题（如果提供）
            if self.document_title:
                title_paragraph = self.document.add_paragraph(self.document_title)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_paragraph.runs[0].font.size = Pt(24)
                title_paragraph.runs[0].font.bold = True
            
            # 解析 HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 处理 HTML 内容
            self._process_html_content(soup)
            
            # 保存文档
            self.document.save(output_file)
            self.logger.info(f"HTML 转换成功，已保存到: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"转换 HTML 时发生错误: {str(e)}")
            return False
    
    def _setup_document_styles(self):
        """设置文档样式"""
        styles = self.document.styles
        
        # 设置正文样式
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.config['font_name']
        normal_font.size = Pt(self.config['font_size'])
        
        # 设置标题样式
        for level, style_config in self.config['heading_styles'].items():
            heading_style_name = f'Heading {level}'
            if heading_style_name in styles:
                heading_style = styles[heading_style_name]
                heading_font = heading_style.font
                heading_font.name = self.config['font_name']
                heading_font.size = Pt(style_config['size'])
                heading_font.bold = style_config['bold']
                heading_font.color.rgb = style_config['color']
    
    def _process_html_content(self, soup: BeautifulSoup):
        """处理 HTML 内容"""
        # 获取 body 内容，如果没有 body 则处理整个文档
        body = soup.find('body')
        if body:
            elements = body.children
        else:
            elements = soup.children
        
        for element in elements:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    self.document.add_paragraph(text)
            elif isinstance(element, Tag):
                self._process_element(element)
    
    def _process_element(self, element: Tag):
        """处理单个 HTML 元素"""
        tag_name = element.name.lower()
        
        # 处理标题
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._add_heading(element)
        # 处理段落
        elif tag_name == 'p':
            self._add_paragraph(element)
        # 处理列表
        elif tag_name in ['ul', 'ol']:
            self._add_list(element)
        # 处理表格
        elif tag_name == 'table':
            self._add_table(element)
        # 处理代码块
        elif tag_name == 'pre':
            self._add_code_block(element)
        # 处理引用
        elif tag_name == 'blockquote':
            self._add_quote(element)
        # 处理图片
        elif tag_name == 'img':
            if not self.config.get('remove_images', False):
                self._add_image(element)
        # 处理链接
        elif tag_name == 'a':
            if not self.config.get('remove_links', False):
                self._add_link(element)
        # 处理 div 和其他容器元素
        elif tag_name in ['div', 'section', 'article', 'main']:
            for child in element.children:
                if isinstance(child, Tag):
                    self._process_element(child)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        self.document.add_paragraph(text)
        # 处理换行
        elif tag_name == 'br':
            pass  # 换行已在段落处理中考虑
        # 处理水平线
        elif tag_name == 'hr':
            self.document.add_paragraph('_' * 50)
        # 其他元素，提取文本内容
        else:
            text = element.get_text().strip()
            if text:
                self.document.add_paragraph(text)
    
    def _add_heading(self, element: Tag):
        """添加标题"""
        tag_name = element.name.lower()
        level = int(tag_name[1])  # h1 -> 1, h2 -> 2, etc.
        title_text = element.get_text().strip()
        
        if title_text:
            heading = self.document.add_heading(title_text, level)
            self._format_paragraph(heading, is_heading=True)
    
    def _add_paragraph(self, element: Tag):
        """添加段落"""
        paragraph = self.document.add_paragraph()
        self._add_inline_content(paragraph, element)
        self._format_paragraph(paragraph)
    
    def _add_inline_content(self, paragraph, element: Tag):
        """添加行内内容（处理粗体、斜体、链接等）"""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    paragraph.add_run(text)
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                
                # 粗体
                if tag_name in ['strong', 'b']:
                    text = child.get_text()
                    run = paragraph.add_run(text)
                    run.bold = True
                # 斜体
                elif tag_name in ['em', 'i']:
                    text = child.get_text()
                    run = paragraph.add_run(text)
                    run.italic = True
                # 链接
                elif tag_name == 'a':
                    if not self.config.get('remove_links', False):
                        text = child.get_text()
                        href = child.get('href', '')
                        run = paragraph.add_run(text)
                        # 可以在这里添加超链接样式
                    else:
                        text = child.get_text()
                        paragraph.add_run(text)
                # 图片
                elif tag_name == 'img':
                    if not self.config.get('remove_images', False):
                        self._add_image(child)
                # 代码
                elif tag_name == 'code':
                    text = child.get_text()
                    run = paragraph.add_run(text)
                    run.font.name = self.config['code_font']
                    run.font.size = Pt(10)
                # 换行
                elif tag_name == 'br':
                    paragraph.add_run('\n')
                # 其他元素，递归处理
                else:
                    self._add_inline_content(paragraph, child)
    
    def _add_list(self, element: Tag):
        """添加列表"""
        tag_name = element.name.lower()
        is_ordered = tag_name == 'ol'
        
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                if is_ordered:
                    paragraph = self.document.add_paragraph(text, style='List Number')
                else:
                    paragraph = self.document.add_paragraph(text, style='List Bullet')
                self._format_paragraph(paragraph)
    
    def _add_table(self, element: Tag):
        """添加表格"""
        rows = element.find_all('tr')
        if not rows:
            return
        
        # 获取列数
        first_row_cells = rows[0].find_all(['td', 'th'])
        cols = len(first_row_cells)
        
        if cols == 0:
            return
        
        # 创建表格
        table = self.document.add_table(rows=len(rows), cols=cols)
        table.style = self.config['table_style']
        
        # 填充表格数据
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < cols:
                    cell_text = cell.get_text().strip()
                    table.cell(row_idx, col_idx).text = cell_text
    
    def _add_code_block(self, element: Tag):
        """添加代码块"""
        code_text = element.get_text()
        
        code_paragraph = self.document.add_paragraph()
        code_run = code_paragraph.add_run(code_text)
        code_run.font.name = self.config['code_font']
        code_run.font.size = Pt(10)
        
        # 设置代码块背景色
        self._set_paragraph_background(code_paragraph)
    
    def _add_quote(self, element: Tag):
        """添加引用块"""
        quote_text = element.get_text().strip()
        
        if quote_text:
            quote_paragraph = self.document.add_paragraph(quote_text)
            quote_paragraph.style = 'Quote'
            self._format_paragraph(quote_paragraph)
    
    def _add_image(self, element: Tag):
        """添加图片"""
        img_url = element.get('src', '')
        alt_text = element.get('alt', '') or element.get('title', '')
        
        if not img_url:
            return
        
        try:
            # 下载图片
            local_path = self._download_image(img_url)
            if not local_path:
                self.logger.warning(f"无法下载图片: {img_url}")
                return
            
            # 处理图片
            converted_path = self._process_image(local_path)
            if not converted_path:
                self.logger.warning(f"图片处理失败: {local_path}")
                return
            
            try:
                with Image.open(converted_path) as img:
                    width, height = img.size
                
                # 计算图片的英寸尺寸
                width_inches = width / 72.0
                height_inches = height / 72.0
                
                # 文档可用宽度约为6.5英寸
                max_width = 6.5
                max_height = 9.0
                
                # 添加图片到文档
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                # 判断图片是否需要缩放
                if width_inches <= max_width and height_inches <= max_height:
                    run.add_picture(converted_path, width=Inches(width_inches), height=Inches(height_inches))
                else:
                    is_landscape = width > height
                    if is_landscape:
                        run.add_picture(converted_path, width=Inches(6))
                    else:
                        run.add_picture(converted_path, height=Inches(6))
                
                self.logger.info(f"添加图片成功: {img_url}")
                
            except Exception as img_error:
                self.logger.error(f"处理图片时出错: {str(img_error)}")
                return
            finally:
                if converted_path != local_path:
                    try:
                        os.remove(converted_path)
                    except Exception:
                        pass
            
            # 添加图片描述
            if alt_text:
                desc_paragraph = self.document.add_paragraph(alt_text)
                desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 删除临时文件
            try:
                os.remove(local_path)
            except Exception:
                pass
                
        except Exception as e:
            self.logger.error(f"添加图片失败: {img_url}, 错误: {str(e)}")
    
    def _add_link(self, element: Tag):
        """添加链接"""
        text = element.get_text().strip()
        if text:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text)
            # 可以在这里添加超链接样式
            self._format_paragraph(paragraph)
    
    def _format_paragraph(self, paragraph, is_heading: bool = False):
        """格式化段落"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = self.config['line_spacing']
        
        if not is_heading:
            paragraph_format.space_before = Pt(self.config['paragraph_spacing_before'])
            paragraph_format.space_after = Pt(self.config['paragraph_spacing_after'])
    
    def _set_paragraph_background(self, paragraph):
        """设置段落背景色（用于代码块）"""
        # 简化实现
        pass
    
    def _process_image(self, image_path: str) -> Optional[str]:
        """处理图片，确保格式兼容性"""
        try:
            with Image.open(image_path) as img:
                original_format = img.format
                
                # 如果是WebP或其他不兼容格式，转换为JPEG
                if original_format in ['WEBP', 'AVIF', 'HEIC', 'HEIF']:
                    temp_dir = tempfile.gettempdir()
                    base_name = os.path.splitext(os.path.basename(image_path))[0]
                    new_path = os.path.join(temp_dir, f"{base_name}_converted.jpg")
                    
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                        img = background
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    img.save(new_path, 'JPEG', quality=85, optimize=True)
                    return new_path
                else:
                    return image_path
                    
        except Exception as e:
            self.logger.error(f"处理图片失败: {image_path}, 错误: {str(e)}")
            return None
    
    def _download_image(self, url: str) -> Optional[str]:
        """下载远程图片到临时文件"""
        try:
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                self.logger.warning(f"无效的图片URL: {url}")
                return None
            
            temp_dir = tempfile.gettempdir()
            filename = sha256(parsed.path.encode()).hexdigest() + ".webp"
            temp_path = os.path.join(temp_dir, filename)
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
            }
            
            response = requests.get(url, stream=True, timeout=20, headers=headers)
            response.raise_for_status()
            
            with open(temp_path, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            
            if not os.path.exists(temp_path):
                return None
            return temp_path
            
        except Exception as e:
            self.logger.error(f"下载图片失败: {url}, 错误: {str(e)}")
            return None


def html_to_docx(html_content: str, output_file: str, 
                document_title: Optional[str] = None,
                config: Optional[Dict[str, Any]] = None) -> bool:
    """
    便捷函数：将 HTML 转换为 Word 文档
    
    Args:
        html_content: HTML 内容
        output_file: 输出的 Word 文件路径
        document_title: 文档标题
        config: 配置选项
        
    Returns:
        是否成功
    """
    converter = HtmlToDocxConverter(config, document_title)
    return converter.convert(html_content, output_file, document_title)


if __name__ == '__main__':
    # 测试示例
    test_html = """
    <h1>标题测试</h1>
    <p>这是一个 <strong>粗体</strong> 和 <em>斜体</em> 的示例。</p>
    <h2>列表示例</h2>
    <ul>
        <li>项目 1</li>
        <li>项目 2</li>
    </ul>
    <h2>代码示例</h2>
    <pre><code>def hello_world():
    print("Hello, World!")</code></pre>
    """
    
    html_to_docx(test_html, 'test_output.docx', '测试文档')
    print("转换完成")
